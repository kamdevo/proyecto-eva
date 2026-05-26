# -*- coding: utf-8 -*-
"""
Genera DOCUMENTACION_TECNICA.docx para el sistema EVA
Estilo: narrativo, con introducción, glosario, texto descriptivo por sección y tablas.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Colores corporativos HUV ──────────────────────────────────────────────────
AZUL_HUV   = RGBColor(0x1D, 0x29, 0x3D)   # azul oscuro corporativo
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)   # azul de sección
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)


# ── Utilidades de formato ─────────────────────────────────────────────────────
def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size or sizes.get(level, 11), bold=True, color=color)
    return p


def add_paragraph(doc, text, indent=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic, color=GRIS_TEXTO)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        cell._tc.tcPr.append(shading)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)

    # Filas
    for ri, row in enumerate(rows):
        bg = 'F2F6FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg)
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(shading)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)

    # Anchos de columna
    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()

    def cp(text, size=12, bold=False, color=GRIS_TEXTO, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
        return p

    cp("HOSPITAL UNIVERSITARIO DEL VALLE", 11, color=AZUL_HUV)
    cp("«Evaristo García» E.S.E.", 10, color=AZUL_HUV)
    doc.add_paragraph()
    cp("DOCUMENTACIÓN TÉCNICA", 22, bold=True, color=AZUL_HUV)
    cp("Sistema EVA", 16, bold=True, color=AZUL_CLARO)
    cp("Sistema de Gestión de Equipos Médicos e Industriales", 12, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Arquitectura, modelo de datos, módulos funcionales y guía técnica", 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Versión 2.0.0  ·  Mayo 2026", 10, color=GRIS_TEXTO)
    cp("Documento confidencial — Prohibida su reproducción sin autorización.", 9, color=RGBColor(0x88, 0x88, 0x88))
    cp("Santiago de Cali, Colombia", 9, color=RGBColor(0x88, 0x88, 0x88))
    doc.add_page_break()


def add_toc_entry(doc, text, page_hint="", level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    indent = Cm(level * 0.6)
    p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    size = 11 if level == 1 else 10
    bold = level == 1
    set_font(run, size=size, bold=bold, color=AZUL_HUV if level == 1 else GRIS_TEXTO)


# ── CONSTRUCCIÓN DEL DOCUMENTO ────────────────────────────────────────────────
def build():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Portada
    add_cover(doc)

    # ── TABLA DE CONTENIDO ────────────────────────────────────────────────────
    add_heading(doc, "Tabla de contenido", 1, AZUL_HUV)
    toc_items = [
        (1, "1.  Introducción"),
        (2, "1.1.  Propósito del sistema"),
        (2, "1.2.  Alcance"),
        (2, "1.3.  Glosario de términos"),
        (1, "2.  Visión general del sistema"),
        (2, "2.1.  Arquitectura en capas"),
        (2, "2.2.  Flujo de la información"),
        (2, "2.3.  Patrón de comunicación frontend-backend"),
        (1, "3.  Stack tecnológico"),
        (2, "3.1.  Backend"),
        (2, "3.2.  Frontend"),
        (1, "4.  Estructura de carpetas"),
        (2, "4.1.  Backend"),
        (2, "4.2.  Frontend"),
        (1, "5.  Módulos del sistema"),
        (2, "5.1.  Módulo Equipos"),
        (2, "5.2.  Módulo Planes de Mantenimiento"),
        (2, "5.3.  Módulo Órdenes / Tickets"),
        (2, "5.4.  Otros módulos"),
        (1, "6.  API REST — Endpoints principales"),
        (2, "6.1.  Autenticación"),
        (2, "6.2.  Equipos"),
        (2, "6.3.  Tickets y correctivos"),
        (2, "6.4.  Planes de mantenimiento"),
        (2, "6.5.  Catálogos"),
        (1, "7.  Sistema de autenticación y permisos"),
        (2, "7.1.  Flujo de inicio de sesión"),
        (2, "7.2.  Roles del sistema"),
        (1, "8.  Generación de reportes y exportaciones"),
        (2, "8.1.  Exportación a Excel (servidor)"),
        (2, "8.2.  Generación de PDF (cliente)"),
        (2, "8.3.  Generación de PDF (servidor)"),
        (1, "9.  Seguridad"),
        (1, "10. Gestión de archivos e imágenes"),
        (1, "11. Caché de opciones de formularios"),
    ]
    for level, text in toc_items:
        add_toc_entry(doc, text, level=level)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCCIÓN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Introducción", 1, AZUL_HUV)
    add_paragraph(doc,
        "El presente documento describe la arquitectura técnica del sistema EVA (Equipos y Vigencias "
        "Asistenciales), la plataforma web diseñada para el Hospital Universitario del Valle «Evaristo García» E.S.E. "
        "con el objetivo de centralizar la gestión del inventario de equipos médicos e industriales, el seguimiento "
        "de su mantenimiento preventivo y correctivo, el control de calibraciones, la trazabilidad de registros INVIMA "
        "y la administración de órdenes de trabajo mediante un sistema de tickets."
    )
    add_paragraph(doc,
        "Este documento está dirigido al equipo de desarrollo, administradores de sistemas e ingenieros biomédicos "
        "que necesiten comprender, mantener o extender las capacidades del sistema. Se cubren los aspectos de "
        "arquitectura, tecnologías empleadas, estructura de código, flujos de datos, módulos funcionales, "
        "endpoints de la API y mecanismos de seguridad."
    )

    add_heading(doc, "1.1. Propósito del sistema", 2, AZUL_CLARO)
    add_paragraph(doc,
        "EVA nació de la necesidad del HUV de reemplazar los procesos manuales y fragmentados de gestión "
        "de equipos biomédicos e industriales. Antes de su implementación, el hospital dependía de hojas de cálculo "
        "desactualizadas, registros en papel y sistemas desconectados que dificultaban el seguimiento del ciclo de "
        "vida de cada equipo, desde su adquisición hasta su baja."
    )
    add_paragraph(doc, "El sistema permite:")
    add_bullet(doc, "Mantener un inventario digital unificado de más de 9.700 equipos biomédicos y 546 equipos industriales.")
    add_bullet(doc, "Registrar y hacer seguimiento de mantenimientos preventivos, correctivos y calibraciones.")
    add_bullet(doc, "Gestionar órdenes de trabajo (tickets) desde su apertura hasta su cierre, con trazabilidad completa.")
    add_bullet(doc, "Controlar la vigencia de registros INVIMA de los dispositivos médicos.")
    add_bullet(doc, "Generar reportes, exportaciones en Excel y fichas técnicas en PDF.")
    add_bullet(doc, "Administrar usuarios, roles y permisos de forma centralizada.")
    add_bullet(doc, "Consultar el historial completo de cada equipo: mantenimientos, correctivos, calibraciones, repuestos.")

    add_heading(doc, "1.2. Alcance", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Este documento cubre la versión 2.0.0 del sistema EVA, desplegada en producción bajo los dominios "
        "eva2.huv.gov.co (frontend) y api.eva2.huv.gov.co (backend). No se incluyen en este documento "
        "los procedimientos de despliegue en servidores (véase la Guía de Despliegue) ni los manuales "
        "de uso para el personal no técnico (véase el Manual de Usuario)."
    )

    add_heading(doc, "1.3. Glosario de términos", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 1. Glosario de términos técnicos usados en este documento.")
    add_table(doc,
        ["Término", "Definición"],
        [
            ["API REST",        "Interfaz de programación que permite la comunicación entre el frontend y el backend mediante HTTP, usando el formato JSON."],
            ["Bearer Token",    "Token de autenticación que el frontend incluye en el encabezado de cada petición HTTP para identificarse ante el backend."],
            ["CORS",            "Cross-Origin Resource Sharing. Mecanismo que controla qué dominios externos pueden hacer peticiones al backend."],
            ["CRUD",            "Create, Read, Update, Delete. Operaciones básicas sobre cualquier entidad del sistema."],
            ["Equipo biomédico","Dispositivo médico de uso clínico o diagnóstico sujeto a regulación INVIMA (Instituto Nacional de Vigilancia de Medicamentos y Alimentos)."],
            ["Equipo industrial","Equipo de infraestructura, planta física, redes o soporte general del hospital."],
            ["Hook",            "Función especial de React que permite usar estado y efectos dentro de componentes funcionales."],
            ["INVIMA",          "Instituto Nacional de Vigilancia de Medicamentos y Alimentos. Entidad colombiana que regula los dispositivos médicos."],
            ["KPI",             "Key Performance Indicator. Indicador clave de rendimiento mostrado en el dashboard."],
            ["Middleware",       "Capa de software que procesa las peticiones HTTP antes de que lleguen al controlador (ej.: autenticación, rate limiting)."],
            ["Modal",           "Ventana emergente dentro de la interfaz web que permite realizar una acción sin abandonar la vista actual."],
            ["OT / Ticket",     "Orden de trabajo de mantenimiento correctivo. Registro digital que documenta la falla, diagnóstico y solución de un equipo."],
            ["Preventivo",      "Mantenimiento programado según el cronograma anual, ejecutado antes de que se produzca una falla."],
            ["Sanctum",         "Paquete de Laravel que gestiona la autenticación de API mediante tokens personales (Bearer Tokens)."],
            ["SPA",             "Single Page Application. La interfaz web se carga una sola vez y las navegaciones posteriores actualizan el contenido sin recargar la página."],
            ["Vite",            "Herramienta de construcción y desarrollo para el frontend. Reemplaza a Webpack con tiempos de compilación significativamente más rápidos."],
        ],
        col_widths=[4.5, 12]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 2. VISIÓN GENERAL
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Visión general del sistema", 1, AZUL_HUV)
    add_paragraph(doc,
        "EVA es una aplicación web que el personal del hospital utiliza desde un navegador, sin necesidad "
        "de instalar ningún software adicional. Sigue una arquitectura cliente-servidor en tres capas: "
        "una interfaz de usuario construida con React (frontend), un servidor de lógica de negocio construido "
        "con Laravel (backend) y una base de datos relacional MySQL. Las tres capas se comunican exclusivamente "
        "mediante una API REST sobre HTTPS, utilizando tokens Bearer para autenticación."
    )

    add_heading(doc, "2.1. Arquitectura en capas", 2, AZUL_CLARO)
    add_paragraph(doc,
        "El sistema se organiza en tres capas claramente diferenciadas. Esta separación facilita el "
        "mantenimiento independiente de cada capa, permite escalar horizontalmente el backend y "
        "hace posible que futuras integraciones (aplicaciones móviles, otros sistemas del hospital) "
        "consuman la misma API sin modificaciones."
    )
    add_paragraph(doc, "Tabla 2. Capas del sistema y sus responsabilidades.")
    add_table(doc,
        ["Capa", "Tecnología", "URL en producción", "Responsabilidad"],
        [
            ["Frontend (cliente)",  "React 19 + Vite + Tailwind CSS v4", "https://eva2.huv.gov.co",      "Interfaz de usuario, navegación, formularios, visualización de datos."],
            ["Backend (servidor)",  "Laravel 12 + PHP 8.2",              "https://api.eva2.huv.gov.co",  "Lógica de negocio, validaciones, autenticación, acceso a base de datos."],
            ["Base de datos",       "MySQL 8.x",                         "Servidor interno HUV",         "Almacenamiento persistente de todos los datos del sistema."],
            ["Almacenamiento",      "Sistema de archivos del servidor",   "storage/app/public/",          "Imágenes, documentos adjuntos, archivos Excel e INVIMA."],
        ],
        col_widths=[3.5, 4, 5, 5]
    )

    add_heading(doc, "2.2. Flujo de la información", 2, AZUL_CLARO)
    add_paragraph(doc,
        "El recorrido de una petición típica desde que el usuario interactúa con la interfaz hasta "
        "que recibe la respuesta es el siguiente:"
    )
    add_bullet(doc, "El usuario realiza una acción en el navegador (ej.: abre el modal de edición de un equipo).")
    add_bullet(doc, "El componente React invoca un hook o servicio que construye una petición HTTP.")
    add_bullet(doc, "Axios envía la petición al backend incluyendo el token Bearer en el encabezado Authorization.")
    add_bullet(doc, "El middleware de Sanctum verifica el token. Si no es válido, devuelve 401 y el frontend redirige al login.")
    add_bullet(doc, "El middleware de permisos verifica que el rol del usuario puede realizar la operación solicitada.")
    add_bullet(doc, "El controlador procesa la petición, ejecuta la lógica de negocio y consulta la base de datos.")
    add_bullet(doc, "La respuesta JSON viaja de regreso al frontend, que actualiza el estado de React y re-renderiza la vista.")

    add_heading(doc, "2.3. Patrón de comunicación frontend-backend", 2, AZUL_CLARO)
    add_paragraph(doc,
        "La comunicación entre el frontend y el backend es completamente stateless: cada petición HTTP "
        "incluye el token de autenticación y toda la información necesaria para ser procesada de forma "
        "independiente. El servidor no mantiene sesiones de usuario en memoria; el estado de sesión "
        "reside exclusivamente en el token gestionado por el cliente."
    )
    add_paragraph(doc,
        "En entorno de desarrollo, Vite actúa como proxy inverso: las peticiones del navegador a /api "
        "son redirigidas automáticamente al servidor Laravel (localhost:8000), evitando problemas de "
        "CORS durante el desarrollo local. En producción, ambos servidores (Nginx para el frontend "
        "y PHP-FPM para el backend) están configurados con los encabezados CORS adecuados."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 3. STACK TECNOLÓGICO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Stack tecnológico", 1, AZUL_HUV)
    add_paragraph(doc,
        "La selección tecnológica de EVA prioriza la estabilidad, el soporte a largo plazo y la "
        "disponibilidad de documentación y comunidad. Todas las tecnologías utilizadas son de código "
        "abierto y cuentan con mantenimiento activo."
    )

    add_heading(doc, "3.1. Backend", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 3. Componentes del backend y su propósito en el sistema.")
    add_table(doc,
        ["Tecnología", "Versión", "Propósito"],
        [
            ["PHP",                        "^8.2",   "Lenguaje de servidor. Se aprovechan las mejoras de rendimiento de PHP 8.x (JIT, fibers, tipos union)."],
            ["Laravel",                    "^12.0",  "Framework principal. Provee el ORM Eloquent, el sistema de rutas, validaciones, colas, eventos y más."],
            ["Laravel Sanctum",            "^4.1",   "Autenticación stateless mediante tokens personales (Bearer Tokens). Reemplaza a Passport para APIs simples."],
            ["Spatie Laravel Permission",  "^6.20",  "Gestión de roles y permisos granulares. Permite definir permisos por módulo y por usuario."],
            ["Maatwebsite Excel",          "^3.1",   "Exportación de datos a formato Excel (.xlsx). Se usa en inventarios, cronogramas y calibraciones."],
            ["barryvdh/laravel-dompdf",   "^3.1",   "Generación de PDFs en el servidor a partir de plantillas Blade."],
            ["MySQL",                      "8.x",    "Base de datos relacional. El esquema tiene más de 100 tablas que cubren todos los módulos del sistema."],
        ],
        col_widths=[4.5, 2.5, 10]
    )

    add_heading(doc, "3.2. Frontend", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 4. Componentes del frontend y su propósito en el sistema.")
    add_table(doc,
        ["Tecnología", "Versión", "Propósito"],
        [
            ["React",                 "19.2.0",    "Biblioteca principal de UI. Se usan componentes funcionales con hooks para todo el estado."],
            ["Vite",                  "^6.3.5",    "Build tool y servidor de desarrollo. Compila el proyecto con esbuild para tiempos de build muy rápidos."],
            ["Tailwind CSS",          "^4.1.10",   "Framework de estilos utilitarios. Permite diseñar directamente en JSX sin archivos CSS separados."],
            ["Shadcn/ui (Radix UI)", "—",          "Componentes de UI accesibles y sin estilos forzados: Dialog, Select, Tabs, Table, etc."],
            ["React Router DOM",      "^7.6.2",    "Enrutamiento del SPA. Define las rutas protegidas y los layouts de la aplicación."],
            ["Axios",                 "^1.10.0",   "Cliente HTTP. Centralizado en httpService.js con interceptores de autenticación y manejo de errores."],
            ["Recharts",              "^3.8.0",    "Biblioteca de gráficos. Se usa en el dashboard para PieChart y BarChart de estadísticas."],
            ["@react-pdf/renderer",   "^4.3.0",    "Generación de PDFs en el cliente (navegador). Se usa para fichas técnicas y tickets."],
            ["Framer Motion",         "^12.x",     "Animaciones de UI: transiciones de modal, efectos de entrada de listas y tooltips."],
            ["Sonner",                "^2.0.6",    "Notificaciones toast (éxito, error, advertencia) accesibles y personalizables."],
            ["Lucide React",          "^0.517.0",  "Biblioteca de íconos SVG. Más de 500 íconos usados en toda la interfaz."],
        ],
        col_widths=[4.5, 2.5, 10]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 4. ESTRUCTURA DE CARPETAS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Estructura de carpetas", 1, AZUL_HUV)
    add_paragraph(doc,
        "El proyecto está organizado en dos repositorios independientes que conviven en el mismo "
        "directorio de trabajo: eva-backend/ para el servidor Laravel y eva-frontend/ para la "
        "aplicación React. Esta separación permite desplegar cada parte de forma independiente "
        "y facilita que equipos diferentes trabajen en paralelo."
    )

    add_heading(doc, "4.1. Backend (eva-backend/)", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 5. Directorios principales del backend y su contenido.")
    add_table(doc,
        ["Directorio / Archivo", "Descripción"],
        [
            ["app/Http/Controllers/Api/", "Controladores de la API REST. Cada módulo tiene su propio controlador."],
            ["app/Http/Middleware/",      "Middleware personalizado: autenticación, permisos, logging de peticiones."],
            ["app/Http/Requests/",        "Clases de validación de formularios (FormRequest). Validan los datos antes de llegar al controlador."],
            ["app/Models/",               "Modelos Eloquent. Representan las tablas de la base de datos y sus relaciones."],
            ["app/Services/",             "Servicios con lógica de negocio reutilizable (ej.: cálculo de fechas de mantenimiento)."],
            ["config/",                   "Configuraciones de Laravel: base de datos, CORS, Sanctum, sesión, correo, etc."],
            ["database/migrations/",      "Definición del esquema de la base de datos en PHP. Permite recrear la estructura en cualquier entorno."],
            ["database/seeders/",         "Datos iniciales del sistema: roles, estados, tipos de equipo, catálogos."],
            ["routes/api.php",            "Todas las rutas de la API REST. Archivo central del sistema, con más de 400 endpoints registrados."],
            ["storage/app/public/",       "Archivos subidos por los usuarios: imágenes de equipos, documentos, archivos Excel, registros INVIMA."],
            [".env",                      "Variables de entorno: credenciales de BD, claves de aplicación, URLs, configuración de correo."],
        ],
        col_widths=[5, 12]
    )

    add_heading(doc, "4.2. Frontend (eva-frontend/src/)", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 6. Directorios del frontend y su contenido.")
    add_table(doc,
        ["Directorio / Archivo", "Descripción"],
        [
            ["App.jsx",              "Componente raíz. Define el router principal, las rutas protegidas y el layout general."],
            ["components/modals/",   "Más de 100 modales de la aplicación: agregar equipo, editar, ver detalle, historial, tickets, etc."],
            ["components/pdf/",      "Generadores de PDF del lado del cliente con @react-pdf/renderer."],
            ["components/ui/",       "Componentes base reutilizables: Button, Input, Select, Table, Badge, Dialog (basados en Shadcn)."],
            ["components/equipment/","Sub-componentes específicos del módulo de equipos."],
            ["components/common/",   "Componentes transversales: Pagination, ConfirmDialog, LoadingSpinner, EmptyState."],
            ["components/skeletons/","Estados de carga (skeleton screens) para cada vista principal."],
            ["components/*.jsx",     "Vistas principales de cada módulo: Equipos, Tickets, Planes, Dashboard, Usuarios, etc."],
            ["contexts/",            "Contextos de React para estado global: AuthContext, ToastContext, TicketsContext, EquipmentSearchContext."],
            ["hooks/",               "Custom hooks: useEquipment, useAuth, useIdleTimeout, useUsuarios, useSedes, etc."],
            ["services/httpService.js", "Instancia de Axios con interceptores de autenticación, manejo de errores y logging."],
            ["services/equipmentPrefetchCache.js", "Caché de opciones de formularios (catálogos) para formularios de equipos."],
            ["config/api.js",        "Constantes globales: URLs de la API, endpoints, timeouts."],
            ["utils/",               "Funciones utilitarias: formateo de fechas, validaciones, manipulación de strings."],
        ],
        col_widths=[5.5, 12]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 5. MÓDULOS DEL SISTEMA
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Módulos del sistema", 1, AZUL_HUV)
    add_paragraph(doc,
        "El sistema EVA está organizado en módulos funcionales accesibles desde el menú lateral "
        "de la aplicación. Cada módulo agrupa un conjunto de funcionalidades relacionadas y tiene "
        "sus propias rutas en el frontend y sus propios endpoints en el backend. El acceso a cada "
        "módulo está controlado por el rol del usuario y por los permisos configurados en la base de datos."
    )

    add_heading(doc, "5.1. Módulo Equipos", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Es el módulo central del sistema. Gestiona el inventario completo de equipos biomédicos "
        "e industriales del hospital. Cada equipo tiene una ficha técnica completa con más de "
        "60 campos: datos de identificación, ubicación, adquisición, mantenimiento, calibración, "
        "INVIMA, imágenes y documentos adjuntos. El módulo incluye búsqueda avanzada, filtros "
        "múltiples, exportación a Excel y generación de ficha técnica en PDF."
    )
    add_paragraph(doc, "Tabla 7. Sub-módulos de Equipos y sus rutas.")
    add_table(doc,
        ["Sub-módulo", "Ruta", "Descripción"],
        [
            ["Biomédicos",      "/equipos/biomedicos",    "Inventario principal de equipos biomédicos. Tabla paginada con búsqueda y filtros por servicio, estado, tipo, etc."],
            ["Industriales",    "/equipos/industriales",  "Inventario de equipos de infraestructura y soporte (plantas, compresores, redes eléctricas, etc.)."],
            ["Órdenes de Compra","/equipos/ordenes-compra","Seguimiento de equipos en proceso de adquisición, con estado de la orden y proveedor."],
            ["Bajas",           "/equipos/bajas",          "Registro de equipos dados de baja, con motivo y documentación de soporte."],
            ["Contingencias",   "/equipos/contingencias",  "Equipos en estado de contingencia o préstamo temporal para reemplazar otro equipo."],
            ["Guías Rápidas",   "/equipos/guias-rapidas",  "Guías de operación rápida para el personal clínico, vinculadas a cada tipo de equipo."],
            ["Manuales",        "/equipos/manuales",       "Repositorio de manuales técnicos de fabricante, descargables desde la plataforma."],
        ],
        col_widths=[3.5, 4.5, 9]
    )

    add_heading(doc, "5.2. Módulo Planes de Mantenimiento Preventivo", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Permite gestionar el cronograma anual de mantenimiento preventivo de los equipos biomédicos. "
        "El coordinador de bioingeniería carga un archivo Excel con la programación del año "
        "(equipo, meses asignados, responsable, frecuencia), y el sistema procesa, almacena y "
        "visualiza el cronograma. Desde la interfaz es posible editar planes individuales, consultar "
        "el historial de cambios de cada plan y hacer seguimiento del cumplimiento."
    )
    add_paragraph(doc, "Tabla 8. Funcionalidades del módulo de Planes de Mantenimiento.")
    add_table(doc,
        ["Funcionalidad", "Descripción"],
        [
            ["Carga de cronograma",     "Importación de archivo Excel con los planes del año. Soporta modo reemplazar o modo agregar."],
            ["Visualización",           "Tabla del cronograma con filtros por año, responsable, frecuencia y estado de cumplimiento."],
            ["Edición de planes",       "Modal para modificar los meses asignados y el responsable de un plan específico."],
            ["Historial de cambios",    "Registro auditable de cada modificación, con usuario y fecha/hora del cambio."],
            ["Registro de ejecución",   "El técnico registra la ejecución del preventivo, adjuntando observaciones y archivo de evidencia."],
            ["Exportación",             "Descarga del cronograma completo en formato Excel."],
        ],
        col_widths=[5, 12]
    )

    add_heading(doc, "5.3. Módulo Órdenes / Tickets", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Sistema de gestión de órdenes de trabajo para mantenimiento correctivo. Cuando un equipo "
        "presenta una falla, el personal clínico o técnico crea un ticket que queda registrado en "
        "el sistema. Los ingenieros biomédicos gestionan el ticket desde su diagnóstico hasta su "
        "cierre, pudiendo registrar avances, adjuntar evidencias y documentar la solución aplicada."
    )
    add_paragraph(doc, "Tabla 9. Vistas del módulo de Tickets y su audiencia.")
    add_table(doc,
        ["Vista", "Ruta", "Audiencia", "Descripción"],
        [
            ["Mis Tickets",        "/ordenes/mis-tickets",       "Todo el personal",   "Tickets creados por el usuario actual, con estado y seguimiento."],
            ["Gestión de Tickets", "/ordenes/gestion-tickets",   "Administradores",    "Panel completo de todos los tickets del sistema, con filtros y estadísticas."],
            ["Tickets Cerrados",   "/ordenes/tickets-cerrados",  "Administradores",    "Historial de tickets finalizados, disponibles para consulta y exportación."],
        ],
        col_widths=[4, 4.5, 3.5, 6]
    )

    add_heading(doc, "5.4. Otros módulos", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 10. Módulos complementarios del sistema.")
    add_table(doc,
        ["Módulo", "Ruta", "Acceso mínimo", "Descripción"],
        [
            ["Calibraciones",    "/calibraciones",        "Rol 3",   "Registro y seguimiento de calibraciones de equipos, con control de vigencia."],
            ["Repuestos",        "/repuestos",            "Rol 3",   "Inventario de repuestos disponibles, con alertas de stock mínimo."],
            ["Capacitaciones",   "/capacitaciones",       "Todos",   "Registro de capacitaciones del personal biomédico por equipo o por servicio."],
            ["Dashboard",        "/dashboard/reportes",   "Rol 1-2", "Panel de indicadores y estadísticas del sistema. KPIs de equipos, correctivos y preventivos."],
            ["Configuración",    "/config/*",             "Rol 1-3", "Administración de catálogos: servicios, áreas, marcas, modelos, estados, etc."],
            ["Administrador",    "/admin/*",              "Rol 1-2", "Gestión de usuarios, roles, permisos y configuración avanzada del sistema."],
            ["Usuarios",         "/admin/usuarios",       "Rol 1-2", "CRUD de usuarios. Búsqueda por nombre, apellido, username o correo electrónico."],
        ],
        col_widths=[3, 4, 3, 8]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 6. API REST
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. API REST — Endpoints principales", 1, AZUL_HUV)
    add_paragraph(doc,
        "Toda la comunicación entre el frontend y el backend se realiza a través de una API REST "
        "con formato JSON. La URL base en producción es https://api.eva2.huv.gov.co/api/v1/. "
        "En desarrollo local, Vite hace proxy de /api hacia http://localhost:8000. "
        "Todas las rutas (salvo el login) requieren el encabezado Authorization: Bearer <token>."
    )

    add_heading(doc, "6.1. Autenticación", 2, AZUL_CLARO)
    add_table(doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["POST", "/auth/login",   "Iniciar sesión. Recibe {email, password}. Devuelve {token, user}."],
            ["POST", "/auth/logout",  "Cerrar sesión. Invalida el token del usuario en la base de datos."],
            ["GET",  "/auth/user",    "Retorna el perfil completo del usuario autenticado."],
        ],
        col_widths=[2, 5, 10]
    )

    add_heading(doc, "6.2. Equipos", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 11. Endpoints del módulo de equipos biomédicos.")
    add_table(doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["GET",    "/equipos",                          "Listado paginado con filtros (search, servicio_id, estado, etc.)."],
            ["GET",    "/equipos/{id}",                     "Detalle completo de un equipo, incluyendo mantenimientos y calibraciones."],
            ["POST",   "/equipos",                          "Crear un nuevo equipo en el inventario."],
            ["PUT",    "/equipos/{id}",                     "Actualizar los datos de un equipo existente."],
            ["DELETE", "/equipos/{id}",                     "Eliminar (baja lógica) un equipo del inventario."],
            ["GET",    "/equipos/medical-devices-complete", "KPIs y métricas globales del inventario para el dashboard."],
            ["GET",    "/equipos/filter-options",           "Catálogos de formularios: marcas, modelos, tipos, estados, servicios, etc."],
            ["POST",   "/equipos/{id}/image",               "Subir o actualizar la imagen principal de un equipo."],
        ],
        col_widths=[2, 6, 9]
    )

    add_heading(doc, "6.3. Tickets y correctivos", 2, AZUL_CLARO)
    add_table(doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["GET",  "/correctivos-generales",       "Listado de tickets correctivos biomédicos, paginado y con filtros."],
            ["POST", "/correctivos-generales",       "Crear un nuevo ticket / orden de trabajo."],
            ["PUT",  "/correctivos-generales/{id}",  "Actualizar estado, diagnóstico o solución de un ticket."],
            ["GET",  "/gestion-tickets",             "Vista administrativa de todos los tickets del sistema."],
        ],
        col_widths=[2, 5.5, 9.5]
    )

    add_heading(doc, "6.4. Planes de mantenimiento", 2, AZUL_CLARO)
    add_table(doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["GET",  "/planes-mantenimientos",                    "Cronograma del año seleccionado con estado de cumplimiento."],
            ["POST", "/planes-mantenimientos/upload-excel",       "Cargar cronograma desde archivo Excel."],
            ["PUT",  "/planes-mantenimientos/{id}",               "Editar meses y responsable de un plan. Registra cambio en auditoría."],
            ["GET",  "/planes-mantenimientos/{id}/historial",     "Historial de cambios de un plan específico con usuario y fecha."],
        ],
        col_widths=[2, 6.5, 8.5]
    )

    add_heading(doc, "6.5. Catálogos de configuración", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 12. Endpoints de catálogos usados en formularios y configuración.")
    add_table(doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["GET", "/servicios",                 "Lista de servicios / unidades del hospital (275 registros)."],
            ["GET", "/areas",                     "Áreas por servicio."],
            ["GET", "/usuarios-public",           "Lista paginada de usuarios. Soporta búsqueda por nombre, apellido, username y correo."],
            ["GET", "/roles",                     "Roles disponibles en el sistema."],
            ["GET", "/sedes",                     "Sedes del hospital."],
            ["GET", "/proveedores-mantenimiento", "Proveedores de mantenimiento registrados."],
            ["GET", "/marcas",                    "Marcas de equipos del catálogo."],
            ["GET", "/modelos",                   "Modelos de equipos del catálogo."],
        ],
        col_widths=[2, 5.5, 9.5]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 7. AUTENTICACIÓN Y PERMISOS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Sistema de autenticación y permisos", 1, AZUL_HUV)
    add_paragraph(doc,
        "EVA utiliza Laravel Sanctum para la autenticación sin estado mediante tokens personales. "
        "Al iniciar sesión, el backend genera un token único que el frontend almacena en memoria "
        "(React Context) y envía en cada petición HTTP. El sistema no usa cookies de sesión, "
        "lo que lo hace compatible con cualquier cliente que pueda hacer peticiones HTTP."
    )

    add_heading(doc, "7.1. Flujo de inicio de sesión", 2, AZUL_CLARO)
    add_bullet(doc, "El usuario ingresa su correo y contraseña en el formulario de login.")
    add_bullet(doc, "El frontend envía POST /api/v1/auth/login con las credenciales.")
    add_bullet(doc, "El backend valida las credenciales contra la tabla 'usuarios' (contraseña Bcrypt, 12 rondas).")
    add_bullet(doc, "Si el correo no está verificado, se retorna error 403 con instrucciones para verificarlo.")
    add_bullet(doc, "Si las credenciales son correctas, se genera un token Sanctum y se devuelve junto con el perfil del usuario.")
    add_bullet(doc, "El frontend almacena el token en AuthContext y redirige al dashboard o a la vista principal.")
    add_bullet(doc, "El hook useIdleTimeout detecta inactividad de 30 minutos y ejecuta el logout automático.")

    add_heading(doc, "7.2. Roles del sistema", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Los permisos se gestionan en dos niveles: a nivel de rol (para el control de acceso a módulos "
        "completos) y a nivel de módulo en la tabla 'permisos' (para control granular de operaciones "
        "CRUD dentro de cada módulo)."
    )
    add_paragraph(doc, "Tabla 13. Roles del sistema y sus privilegios.")
    add_table(doc,
        ["rol_id", "Nombre", "Privilegios principales"],
        [
            ["1", "Super Administrador", "Acceso total al sistema, incluyendo administración de otros administradores y configuración avanzada."],
            ["2", "Administrador",       "Acceso completo a todos los módulos operativos. Sin acceso a gestión de super-administradores."],
            ["3", "Avanzado",            "Acceso a equipos, tickets, mantenimientos y calibraciones. Sin acceso a usuarios ni configuración avanzada."],
            ["4", "Normal",              "Solo lectura del inventario de equipos y gestión de sus propios tickets. Sin exportaciones ni configuración."],
        ],
        col_widths=[2, 4, 11]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 8. REPORTES Y EXPORTACIONES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Generación de reportes y exportaciones", 1, AZUL_HUV)
    add_paragraph(doc,
        "El sistema ofrece múltiples mecanismos para extraer información en formatos estándar. "
        "Las exportaciones se pueden generar desde cualquier módulo principal y no requieren "
        "configuración adicional; el archivo se descarga directamente en el navegador del usuario."
    )

    add_heading(doc, "8.1. Exportación a Excel (servidor)", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Utiliza el paquete Maatwebsite\\Excel en el backend. El servidor genera el archivo .xlsx "
        "y lo devuelve como respuesta HTTP con Content-Type adecuado. El frontend recibe el "
        "contenido como un Blob binario y crea un enlace de descarga temporal en el navegador, "
        "sin abrir una nueva pestaña."
    )
    add_paragraph(doc, "Exportaciones disponibles:")
    add_bullet(doc, "Inventario de equipos biomédicos (con todos los campos de la ficha técnica).")
    add_bullet(doc, "Inventario de equipos industriales.")
    add_bullet(doc, "Cronograma de mantenimiento preventivo anual.")
    add_bullet(doc, "Historial de mantenimientos preventivos ejecutados.")
    add_bullet(doc, "Registros de calibraciones (biomédicas e industriales).")
    add_bullet(doc, "Tickets y correctivos de mantenimiento.")

    add_heading(doc, "8.2. Generación de PDF (cliente)", 2, AZUL_CLARO)
    add_paragraph(doc,
        "Utiliza la biblioteca @react-pdf/renderer en el frontend. El PDF se genera directamente "
        "en el navegador del usuario sin necesidad de peticiones adicionales al servidor. "
        "Esta aproximación reduce la carga del servidor y permite generación instantánea."
    )
    add_paragraph(doc,
        "Nota técnica importante: @react-pdf/renderer no puede renderizar HTML. Toda descripción "
        "que provenga del sistema con formato HTML (negritas, listas, cursivas generadas por el "
        "editor de texto enriquecido) debe pasar por la función stripHtml() antes de incluirse "
        "en el PDF, para eliminar las etiquetas y mostrar solo el texto plano."
    )
    add_paragraph(doc, "Componentes PDF disponibles:")
    add_bullet(doc, "TicketPDF.jsx — Ficha completa de una orden de trabajo / ticket de correctivo.")
    add_bullet(doc, "equipment-modal-replica-pdf.jsx — Ficha técnica completa de un equipo biomédico.")

    add_heading(doc, "8.3. Generación de PDF (servidor)", 2, AZUL_CLARO)
    add_paragraph(doc,
        "El paquete barryvdh/laravel-dompdf se usa en el backend para reportes que requieren "
        "plantillas Blade de Laravel con formato HTML complejo, como informes periódicos y "
        "documentos con membrete institucional del hospital."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 9. SEGURIDAD
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Seguridad", 1, AZUL_HUV)
    add_paragraph(doc,
        "El sistema EVA implementa múltiples capas de seguridad que cubren desde el transporte "
        "de datos hasta la autorización granular de operaciones. A continuación se describen "
        "los mecanismos principales."
    )
    add_paragraph(doc, "Tabla 14. Mecanismos de seguridad implementados en EVA.")
    add_table(doc,
        ["Mecanismo", "Implementación", "Descripción"],
        [
            ["Autenticación",       "Laravel Sanctum (Bearer Tokens)",      "Cada petición debe incluir un token válido. Los tokens se almacenan hasheados en la BD."],
            ["Autorización",        "Spatie Permission + permisos propios",  "Dos capas: control por rol (módulos completos) y permisos granulares CRUD por módulo."],
            ["Rate Limiting",       "throttle:60,1 en todas las rutas API",  "Máximo 60 peticiones por minuto por IP. Protege contra ataques de fuerza bruta."],
            ["CORS",                "config/cors.php con lista blanca",      "Solo los dominios autorizados (HUV production + localhost dev) pueden acceder a la API."],
            ["Validación",          "FormRequest en cada endpoint",          "Los datos de entrada se validan estrictamente antes de procesarse. Se rechazan peticiones malformadas."],
            ["Contraseñas",         "Bcrypt (12 rondas)",                    "Las contraseñas nunca se almacenan en texto plano. El costo de 12 rondas dificulta ataques de diccionario."],
            ["Sesión inactiva",     "useIdleTimeout hook (frontend)",        "Logout automático tras 30 minutos de inactividad del usuario en el navegador."],
            ["Verificación email",  "Verificación obligatoria",              "Las cuentas nuevas deben verificar su correo antes de poder acceder al sistema."],
            ["HTTPS",               "Certificado SSL en producción",         "Todo el tráfico entre el navegador y los servidores viaja cifrado."],
        ],
        col_widths=[3.5, 4.5, 9]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 10. GESTIÓN DE ARCHIVOS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Gestión de archivos e imágenes", 1, AZUL_HUV)
    add_paragraph(doc,
        "El sistema gestiona diferentes tipos de archivos adjuntos asociados a equipos y registros. "
        "Todos los archivos se almacenan en el sistema de archivos del servidor bajo el directorio "
        "storage/app/public/, y se sirven a través de endpoints específicos que verifican "
        "la autenticación antes de entregar el archivo."
    )
    add_paragraph(doc, "Tabla 15. Tipos de archivos y sus rutas de almacenamiento.")
    add_table(doc,
        ["Tipo de archivo", "Ruta física en servidor", "Endpoint de descarga"],
        [
            ["Imágenes de equipos",       "storage/app/public/equipos/images/",  "GET /api/v1/equipos/{id}/image/{filename}"],
            ["Documentos de equipos",     "storage/app/public/equipos/archivos/", "GET /api/v1/equipos/download/{filename}"],
            ["Archivos de correctivos",   "storage/app/public/correctivos/",      "GET /api/v1/download/correctivos/{filename}"],
            ["Registros INVIMA",          "storage/app/public/invimas/",          "GET /api/v1/invima/download/{filename}"],
            ["Archivos de preventivos",   "storage/app/public/mantenimientos/",   "GET /api/v1/download/mantenimientos/{filename}"],
            ["Observaciones",             "storage/app/public/observaciones/",    "GET /api/v1/download/observaciones/{filename}"],
            ["Repuestos",                 "storage/app/public/repuestos/",        "GET /api/v1/download/repuestos/{filename}"],
            ["Órdenes de compra",         "storage/app/public/ordenes_compra/",   "GET /api/v1/download/ordenes-compra/{filename}"],
        ],
        col_widths=[4.5, 5.5, 7]
    )
    add_paragraph(doc,
        "Las imágenes de equipos se sirven con carga progresiva en el frontend: el componente "
        "EquipmentImage.jsx muestra primero un skeleton placeholder mientras la imagen carga, "
        "y en caso de error de carga muestra una imagen por defecto. Esto mejora la percepción "
        "de velocidad en listados con muchos equipos."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 11. CACHÉ DE FORMULARIOS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. Caché de opciones de formularios", 1, AZUL_HUV)
    add_paragraph(doc,
        "Los formularios de creación y edición de equipos incluyen numerosos campos de selección "
        "(marcas, modelos, servicios, áreas, estados, propietarios, frecuencias, períodos de garantía, "
        "etc.). Cargar todos estos catálogos desde el servidor cada vez que el usuario abre un "
        "formulario resultaría en múltiples peticiones HTTP y una experiencia lenta."
    )
    add_paragraph(doc,
        "Para resolver esto, el servicio equipmentPrefetchCache.js implementa un sistema de caché "
        "en memoria. Cuando el usuario abre por primera vez el modal de agregar o editar equipo, "
        "se realiza una única petición a GET /equipos/filter-options que devuelve todos los "
        "catálogos de una vez. El resultado se almacena en memoria durante la sesión y se reutiliza "
        "en aperturas posteriores del formulario. El caché se invalida automáticamente cuando el "
        "usuario crea o edita un registro que podría afectar los catálogos."
    )
    add_paragraph(doc, "Catálogos incluidos en el caché:")
    add_bullet(doc, "Tipos de equipo, marcas, modelos y clasificaciones biomédicas.")
    add_bullet(doc, "Servicios hospitalarios y áreas por servicio.")
    add_bullet(doc, "Estados de equipo, propietarios, disponibilidades.")
    add_bullet(doc, "Períodos de garantía, frecuencias de mantenimiento.")
    add_bullet(doc, "Fuentes de alimentación eléctrica, tecnologías predominantes.")
    add_bullet(doc, "Tipos de adquisición y riesgos según clasificación biomédica.")

    # Guardar
    output_path = os.path.join(OUTPUT_DIR, "DOCUMENTACION_TECNICA.docx")
    doc.save(output_path)
    print(f"Generado: {output_path}")


if __name__ == "__main__":
    build()
