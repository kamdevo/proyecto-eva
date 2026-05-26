# -*- coding: utf-8 -*-
"""Genera MANUAL_DE_USUARIO.docx para el sistema EVA"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
AZUL_HUV   = RGBColor(0x1D, 0x29, 0x3D)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name, run.font.size, run.font.bold, run.font.italic = 'Calibri', Pt(size), bold, italic
    if color: run.font.color.rgb = color

def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size or sizes.get(level, 11), bold=True, color=color)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style, table.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        cell._tc.get_or_add_tcPr().append(shading)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F6FB' if ri % 2 == 0 else 'FFFFFF')
            cell._tc.get_or_add_tcPr().append(shading)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)
    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)
    doc.add_paragraph()

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin, section.right_margin = Cm(3), Cm(2.5)
    
    for _ in range(6): doc.add_paragraph()
    def cp(text, size=12, bold=False, color=GRIS_TEXTO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    cp("HOSPITAL UNIVERSITARIO DEL VALLE", 11, color=AZUL_HUV)
    cp("«Evaristo García» E.S.E.", 10, color=AZUL_HUV)
    doc.add_paragraph()
    cp("MANUAL DE USUARIO", 22, bold=True, color=AZUL_HUV)
    cp("Sistema EVA — Guía completa para usuarios finales", 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Versión 2.0.0  ·  Mayo 2026", 10, color=GRIS_TEXTO)
    doc.add_page_break()

    add_heading(doc, "1. Bienvenida")
    add_paragraph(doc, "Este manual está diseñado para el personal clínico, técnico y administrativo del Hospital Universitario del Valle que utiliza el sistema EVA. Le guiará paso a paso en el uso de cada módulo, desde el inicio de sesión hasta la generación de reportes.")
    add_paragraph(doc, "Si tiene preguntas o encuentra problemas, contacte al departamento de Bioingeniería o escriba un ticket de soporte dentro de la aplicación.")

    add_heading(doc, "2. Inicio de sesión")
    add_paragraph(doc, "Acceda a la aplicación ingresando su correo electrónico y contraseña en eva2.huv.gov.co. Si olvida su contraseña, haga clic en 'Recuperar contraseña' e ingrese su correo para recibir un enlace de restablecimiento.")
    add_bullet(doc, "Usuario: correo electrónico registrado en el sistema.")
    add_bullet(doc, "Contraseña: contraseña personal asignada por el administrador (o restablecida por usted).")
    add_bullet(doc, "La sesión expira automáticamente después de 30 minutos de inactividad.")

    add_heading(doc, "3. Módulo de Equipos")
    add_heading(doc, "3.1. Consultar un equipo biomédico", 2, AZUL_CLARO)
    add_paragraph(doc, "1. En el menú lateral, haga clic en 'Equipos' → 'Biomédicos'. 2. Busque por nombre, código o servicio. 3. Haga clic en la fila del equipo para ver todos sus detalles, incluyendo mantenimientos previos, calibraciones y documentos adjuntos.")
    
    add_heading(doc, "3.2. Descargar ficha técnica en PDF", 2, AZUL_CLARO)
    add_paragraph(doc, "Abra el modal de detalle del equipo. En la esquina superior derecha encontrará el botón 'Descargar PDF'. El navegador descargará un archivo que contiene toda la información del equipo de forma imprimible.")

    add_heading(doc, "4. Módulo de Tickets")
    add_heading(doc, "4.1. Crear un nuevo ticket", 2, AZUL_CLARO)
    add_paragraph(doc, "1. Haga clic en 'Órdenes' → 'Mis Tickets'. 2. Haga clic en 'Crear Ticket'. 3. Seleccione el equipo afectado. 4. Describa el problema con el mayor detalle posible. 5. Adjunte fotos o documentos si es necesario. 6. Haga clic en 'Crear'.")

    add_heading(doc, "4.2. Hacer seguimiento del ticket", 2, AZUL_CLARO)
    add_paragraph(doc, "En la vista 'Mis Tickets', verá una lista de todos sus tickets con su estado actual (Abierto, En proceso, Resuelto, Cerrado). Haga clic en cualquier ticket para ver el historial de avances y cambios.")

    add_heading(doc, "5. Módulo de Mantenimiento Preventivo")
    add_heading(doc, "5.1. Consultar el cronograma", 2, AZUL_CLARO)
    add_paragraph(doc, "Haga clic en 'Planes' → 'Preventivo'. Seleccione el año del cronograma (2024, 2025, etc.). Verá una tabla con todos los equipos programados, los meses asignados y el responsable del mantenimiento.")

    add_heading(doc, "5.2. Registrar ejecución de preventivo", 2, AZUL_CLARO)
    add_paragraph(doc, "1. Identifique el plan en la tabla. 2. Haga clic en 'Registrar ejecución'. 3. Adjunte evidencias (fotos, reportes). 4. Agregue observaciones si es necesario. 5. Haga clic en 'Guardar'. El sistema actualizará el estado de cumplimiento automáticamente.")

    add_heading(doc, "6. Exportación de datos")
    add_paragraph(doc, "En cualquier tabla (equipos, tickets, cronograma), encontrará un botón 'Descargar Excel' que genera un archivo con toda la información visible, listo para analizar en Microsoft Excel o Google Sheets.")
    
    add_heading(doc, "7. Buscar un usuario")
    add_paragraph(doc, "En la página de administración de usuarios, utilice el campo de búsqueda para encontrar usuarios por nombre, apellido, usuario o correo electrónico. El sistema filtrará los resultados en tiempo real.")

    add_heading(doc, "8. Recomendaciones de uso")
    add_bullet(doc, "Cierre sesión cuando termine de usar la aplicación, especialmente en equipos compartidos.")
    add_bullet(doc, "No comparta su contraseña con otros usuarios.")
    add_bullet(doc, "Use descripciones claras y detalladas al crear tickets o registrar problemas.")
    add_bullet(doc, "Adjunte fotografías de alto contraste para evidencias de mantenimiento.")
    add_bullet(doc, "Revise regularmente sus tickets pendientes y actualice su estado según sea necesario.")

    output_path = os.path.join(OUTPUT_DIR, "MANUAL_DE_USUARIO.docx")
    doc.save(output_path)
    print(f"Generado: {output_path}")

if __name__ == "__main__":
    build()
