# -*- coding: utf-8 -*-
"""Genera Estadistica.docx con mayor contenido narrativo para el sistema EVA"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
AZUL_HUV = RGBColor(0x1D, 0x29, 0x3D)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
VERDE = RGBColor(0x2E, 0x7D, 0x32)


def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=size or sizes.get(level, 11), bold=True, color=color)


def add_paragraph(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text)
    set_font(r, size=10, color=GRIS_TEXTO)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=10, color=GRIS_TEXTO)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        cell._tc.get_or_add_tcPr().append(shading)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F6FB' if ri % 2 == 0 else 'FFFFFF')
            cell._tc.get_or_add_tcPr().append(shading)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)

    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)

    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Portada
    for _ in range(6):
        doc.add_paragraph()

    def center_line(text, size=12, bold=False, color=GRIS_TEXTO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)

    center_line('HOSPITAL UNIVERSITARIO DEL VALLE', 11, color=AZUL_HUV)
    center_line('«Evaristo García» E.S.E.', 10, color=AZUL_HUV)
    doc.add_paragraph()
    center_line('ESTADÍSTICAS DEL SISTEMA EVA', 22, bold=True, color=AZUL_HUV)
    center_line('Cifras reales de la base de datos — Hospital Universitario del Valle', 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    center_line('Versión 2.0.1  |  Mayo 2026', 10, color=GRIS_TEXTO)
    doc.add_page_break()

    add_heading(doc, '1. Resumen ejecutivo')
    add_paragraph(doc, 'Este documento consolida las principales cifras operativas del sistema EVA a partir de la base de datos local gestionthuv. Su objetivo es ofrecer una lectura rápida del tamaño real del inventario, el volumen de mantenimientos ejecutados, el comportamiento de las calibraciones y la cobertura funcional del dashboard.')
    add_paragraph(doc, 'La estadística no solo presenta números: también ayuda a interpretar la carga de trabajo de bioingeniería, el alcance de la plataforma y el uso esperado de cada módulo. Por eso, cada bloque incluye un breve contexto para que la información sea útil tanto para toma de decisiones como para auditoría interna.')
    add_bullet(doc, 'El sistema administra miles de equipos biomédicos e industriales con trazabilidad completa.')
    add_bullet(doc, 'Los mantenimientos preventivos concentran el mayor volumen histórico de registros.')
    add_bullet(doc, 'Las exportaciones y el dashboard permiten consultar indicadores sin acceder directamente a la base de datos.')
    add_bullet(doc, 'Los datos visibles aquí corresponden a la consulta local utilizada para validar la estructura del informe.')

    add_heading(doc, '2. Metodología de consulta', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'Las cifras incluidas en este documento se obtuvieron mediante consultas agregadas sobre las tablas principales del sistema. Se priorizó el conteo de registros, la clasificación por módulo y la identificación de indicadores que ya están expuestos en la interfaz del frontend.')
    add_paragraph(doc, 'En términos prácticos, esta estadística responde a tres preguntas: cuántos activos tiene el hospital, qué tan grande es la operación de mantenimiento y qué tan preparada está la plataforma para exportar y visualizar información de forma confiable.')
    add_table(doc, ['Fuente', 'Uso en este documento', 'Observación'], [
        ['Base de datos local', 'Conteo real de equipos, tickets, planes, calibraciones y usuarios', 'gestionthuv'],
        ['Frontend React', 'KPI mostrados en el dashboard y módulos visibles al usuario', 'Se toma como referencia funcional'],
        ['Backend Laravel', 'Exportaciones y endpoints de consulta', 'Información descargable en Excel'],
    ], col_widths=[4.5, 7.5, 5.5])

    add_heading(doc, '3. Panorama general del sistema', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'El siguiente bloque resume el inventario de datos más importante para entender el alcance de EVA. La lectura combinada de estas cifras muestra que el sistema no solo sirve como catálogo de equipos, sino también como repositorio histórico de operación y control de mantenimiento.')
    add_table(doc, ['Módulo', 'Tabla BD', 'Cantidad total'], [
        ['Equipos biomédicos', 'equipos', '9.765'],
        ['Equipos industriales', 'equipos_industriales', '546'],
        ['Usuarios registrados', 'usuarios', '269'],
        ['Servicios / Unidades', 'servicios', '275'],
        ['Centros de costo', 'centros', '334'],
        ['Ubicaciones', 'ubicaciones', '244'],
        ['Marcas', 'marcas', '1.013'],
        ['Modelos', 'modelos', '2.468'],
        ['Repositorios documentales', 'documentos', '1.119'],
        ['Tickets / Correctivos', 'correctivos_generales + correctivos_generales_ind', '2.435'],
        ['Preventivos ejecutados', 'mantenimiento + mantenimiento_ind', '16.859'],
        ['Calibraciones registradas', 'calibracion + calibracion_ind', '8.897'],
    ], col_widths=[5.2, 5.5, 5.3])
    add_paragraph(doc, 'La magnitud de estos números confirma que EVA tiene un uso institucional sostenido. En particular, el volumen de equipos biomédicos y la cantidad de registros históricos de preventivos evidencian una operación madura, donde el mantenimiento preventivo es el componente más representativo en términos de seguimiento y control.')

    add_heading(doc, '4. Equipos biomédicos', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'El inventario biomédico constituye la base operativa del sistema. Aquí se concentra la mayor cantidad de recursos, así como la trazabilidad de mantenimiento, calibración, estado y ubicación de cada dispositivo.')
    add_paragraph(doc, 'Una proporción alta de equipos activos indica que la plataforma se usa para monitorear infraestructura real en servicio. Las pequeñas diferencias en estados no activos suelen corresponder a baja, préstamo, reparación o procesos de depuración de inventario.')
    add_table(doc, ['Indicador', 'Valor'], [
        ['Total equipos en inventario', '9.765'],
        ['Equipos en estado Activo', '9.762'],
        ['Equipos con otros estados', '3'],
        ['Estados disponibles en catálogo', '16 (Activo, Baja, Préstamo, En reparación, etc.)'],
    ], col_widths=[7.2, 10])
    add_paragraph(doc, 'Top 10 servicios con más equipos biomédicos: este análisis permite visualizar qué áreas concentran la mayor carga de activos y, por tanto, requieren mayor prioridad de soporte técnico y planificación de mantenimiento.')
    add_table(doc, ['#', 'Servicio / Unidad', 'Equipos'], [
        ['1', 'URGENCIAS', '689'],
        ['2', 'QUIRÓFANO CENTRAL', '412'],
        ['3', 'TALLER BIOMÉDICO NORTE', '398'],
        ['4', 'HOSPITALIZACIÓN CIRUGÍA', '357'],
        ['5', 'LABORATORIO', '336'],
        ['6', 'UCI ADULTOS', '321'],
        ['7', 'UCI PEDIÁTRICA', '287'],
        ['8', 'IMAGENOLOGÍA', '269'],
        ['9', 'HOSPITALIZACIÓN MEDICINA INTERNA', '254'],
        ['10', 'NEONATOS', '238'],
    ], col_widths=[1.2, 9.6, 2.5])

    add_heading(doc, '5. Mantenimiento correctivo', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'Los correctivos representan la respuesta operativa a fallas reportadas por los usuarios o detectadas durante inspecciones. Su volumen sirve como termómetro de incidentes, desgaste de equipos y demanda de atención técnica.')
    add_paragraph(doc, 'Un número elevado de correctivos no necesariamente indica un problema del sistema; también puede reflejar mayor cobertura de uso, más capacidad de registro y una cultura institucional de reporte oportuno. Lo importante es que cada caso quede documentado con diagnóstico, solución y cierre.')
    add_table(doc, ['Indicador', 'Tabla BD', 'Valor'], [
        ['Correctivos biomédicos registrados', 'correctivos_generales', '2.282'],
        ['Correctivos industriales registrados', 'correctivos_generales_ind', '153'],
        ['Total correctivos en el sistema', 'ambas tablas', '2.435'],
    ], col_widths=[6.5, 5.5, 4.5])
    add_paragraph(doc, 'El predominio de los correctivos biomédicos es coherente con la composición del inventario: el sistema administra más equipos clínicos que industriales y, por tanto, el flujo de incidencias también se concentra allí.')

    add_heading(doc, '6. Mantenimiento preventivo', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'Este es el bloque histórico más voluminoso del informe. Los preventivos permiten verificar el cumplimiento del cronograma de mantenimiento, medir productividad técnica y evaluar si el hospital mantiene sus activos bajo una política preventiva consistente.')
    add_paragraph(doc, 'Además de los registros ejecutados, el sistema conserva planes cargados por año. Esa información es fundamental para reconstruir la evolución del programa de mantenimiento y detectar períodos donde la carga operativa fue más alta.')
    add_table(doc, ['Indicador', 'Tabla BD', 'Valor'], [
        ['Preventivos biomédicos ejecutados', 'mantenimiento', '16.842'],
        ['Preventivos industriales ejecutados', 'mantenimiento_ind', '17'],
        ['Total preventivos ejecutados', 'ambas tablas', '16.859'],
        ['Planes de mantenimiento cargados', 'planes_mantenimientos', '6.592'],
    ], col_widths=[6.5, 5.5, 4.5])
    add_paragraph(doc, 'Plan de mantenimiento por año: esta vista ayuda a identificar en qué periodos se consolidaron más registros y cómo se distribuye el histórico entre vigencias.')
    add_table(doc, ['Año', 'Planes cargados', 'Tabla BD'], [
        ['2025', '1', 'planes_mantenimientos'],
        ['2024', '891', 'planes_mantenimientos'],
        ['2022', '1.644', 'planes_mantenimientos'],
        ['2021', '1.661', 'planes_mantenimientos'],
        ['2020', '1.310', 'planes_mantenimientos'],
        ['2019', '912', 'planes_mantenimientos'],
        ['2018', '173', 'planes_mantenimientos'],
    ], col_widths=[2.5, 4.3, 7.8])

    add_heading(doc, '7. Calibraciones', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'Las calibraciones son un componente crítico para equipos de medición, diagnóstico y soporte clínico. El sistema distingue registros biomédicos e industriales para mantener trazabilidad y cumplimiento normativo.')
    add_paragraph(doc, 'Este bloque es especialmente importante para auditorías internas y seguimiento de vigencia, porque permite comprobar que los dispositivos sensibles se han revisado con la periodicidad requerida.')
    add_table(doc, ['Indicador', 'Tabla BD', 'Valor'], [
        ['Calibraciones biomédicas registradas', 'calibracion', '8.576'],
        ['Calibraciones industriales registradas', 'calibracion_ind', '321'],
        ['Total calibraciones en el sistema', 'ambas tablas', '8.897'],
        ['Equipos con calibración visible en reportes', 'front + backend', 'Cobertura completa en el módulo'],
    ], col_widths=[6.5, 5.5, 4.5])

    add_heading(doc, '8. KPIs del dashboard', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'El dashboard expone una selección de indicadores pensados para lectura inmediata por administradores y coordinadores. Aquí se priorizan métricas operativas que permiten detectar volumen de inventario, carga de mantenimiento y avance de exportaciones sin entrar al detalle de cada tabla.')
    add_paragraph(doc, 'En la práctica, estos KPIs sirven como capa ejecutiva del sistema: resumen en una sola vista la situación general de equipos, correctivos y planes activos, y facilitan la comunicación entre las áreas técnica y administrativa.')
    add_table(doc, ['KPI', 'Valor actual (local)', 'Fuente (endpoint API)', 'Tabla BD'], [
        ['Total equipos biomédicos', '9.765', 'GET /equipos/medical-devices-complete', 'equipos'],
        ['Equipos activos', '9.762', 'GET /equipos/medical-devices-complete', 'equipos'],
        ['Equipos industriales', '546', 'GET /equipos-industriales', 'equipos_industriales'],
        ['Correctivos biomédicos', '2.282', 'GET /correctivos-generales', 'correctivos_generales'],
        ['Correctivos industriales', '153', 'GET /correctivos-generales-ind', 'correctivos_generales_ind'],
        ['Preventivos biomédicos', '16.842', 'GET /mantenimiento/export', 'mantenimiento'],
        ['Preventivos industriales', '17', 'GET /mantenimiento-ind/export', 'mantenimiento_ind'],
        ['Calibraciones biomédicas', '8.576', 'GET /calibraciones/export', 'calibracion'],
        ['Calibraciones industriales', '321', 'GET /calibraciones-ind/export', 'calibracion_ind'],
        ['Planes de mantenimiento', '6.592', 'GET /planes-mantenimientos/export', 'planes_mantenimientos'],
        ['Usuarios registrados', '269', 'GET /usuarios', 'usuarios'],
        ['Servicios / Unidades', '275', 'GET /servicios', 'servicios'],
    ], col_widths=[4.4, 3.5, 6.8, 4.1])

    add_heading(doc, '9. Exportaciones disponibles', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'Una de las ventajas más prácticas de EVA es que convierte las consultas operativas en archivos descargables para análisis externo. Esto reduce la dependencia de consultas manuales a base de datos y acelera la elaboración de informes de gestión.')
    add_paragraph(doc, 'Las exportaciones están pensadas para usuarios administrativos y técnicos que necesitan consolidar información en Excel, filtrar por año o revisar históricos de forma offline.')
    add_table(doc, ['Exportación', 'Formato', 'Registros aprox.', 'Endpoint Backend'], [
        ['Correctivos biomédicos', 'Excel (.xlsx)', '2.282', 'GET /correctivos-generales/export'],
        ['Correctivos industriales', 'Excel (.xlsx)', '153', 'GET /correctivos-generales-ind/export'],
        ['Preventivos biomédicos', 'Excel (.xlsx)', '16.842', 'GET /mantenimiento/export'],
        ['Calibraciones', 'Excel (.xlsx)', '8.897', 'GET /calibraciones/export'],
        ['Cronograma de mto.', 'Excel (.xlsx)', '6.592', 'GET /planes-mantenimientos/export'],
        ['Inventario de equipos', 'Excel (.xlsx)', '10.311', 'GET /equipos/export'],
    ], col_widths=[5.1, 3.2, 3.2, 7.0])
    add_paragraph(doc, 'Las exportaciones implementadas con Maatwebsite/Excel se descargan como archivo local, lo que evita que el usuario tenga que abrir otra pestaña o copiar datos manualmente.')

    add_heading(doc, '10. Acceso al dashboard por rol', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'El acceso al dashboard no es uniforme para todos los perfiles. La visibilidad de indicadores y la capacidad de exportación se regulan mediante roles, de modo que cada usuario vea solo lo que corresponde a su nivel de responsabilidad.')
    add_paragraph(doc, 'Esto ayuda a mantener el orden operativo y a evitar que información sensible se exponga a usuarios que solo requieren lectura básica.')
    add_table(doc, ['Rol', 'Dashboard', 'Exportaciones', 'Gestión equipos'], [
        ['Super Administrador (1)', 'Acceso completo', 'Todas', 'Sí'],
        ['Administrador (2)', 'Acceso completo', 'Todas', 'Sí'],
        ['Avanzado (3)', 'No visible', '—', 'Limitado'],
        ['Normal (4)', 'No visible', '—', 'Solo lectura'],
    ], col_widths=[5.5, 4.5, 3.5, 4.5])

    add_heading(doc, '11. Glosario', level=2, color=AZUL_CLARO)
    add_paragraph(doc, 'Para facilitar la lectura del informe, se incluyen algunos términos clave que aparecen con frecuencia en el sistema y en las consultas técnicas.')
    add_table(doc, ['Término', 'Definición'], [
        ['KPI', 'Key Performance Indicator — Indicador clave de rendimiento'],
        ['OT / Correctivo', 'Orden de trabajo de mantenimiento correctivo'],
        ['Vigencia activa', 'Año del plan de mantenimiento definido en la tabla vigencias_mantenimiento'],
        ['Biomédico', 'Equipo de uso clínico o diagnóstico bajo regulación INVIMA'],
        ['Industrial', 'Equipo de infraestructura, planta física, redes o soporte general'],
        ['Preventivo', 'Mantenimiento programado para evitar fallas y extender vida útil'],
        ['Calibración', 'Proceso de verificación y ajuste de precisión del equipo'],
        ['Exportación', 'Archivo descargable generado por el backend en formato Excel'],
    ], col_widths=[4.2, 13.0])

    add_heading(doc, '12. Conclusión')
    add_paragraph(doc, 'En conjunto, las cifras confirman que EVA es una plataforma de alto uso institucional, con una base sólida de activos, una carga histórica significativa de mantenimiento y mecanismos suficientes para reportar, exportar y supervisar la operación de manera centralizada.')
    add_paragraph(doc, 'La lectura de estas estadísticas debe entenderse como una referencia operativa y no solo como un inventario contable. Cada número refleja actividad real de mantenimiento, trazabilidad técnica y el nivel de madurez con el que el hospital administra sus recursos biomédicos e industriales.')

    output_path = os.path.join(OUTPUT_DIR, 'Estadistica.docx')
    doc.save(output_path)
    print(f'Generado: {output_path}')


if __name__ == '__main__':
    build()
