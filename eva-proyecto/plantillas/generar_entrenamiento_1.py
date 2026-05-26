# -*- coding: utf-8 -*-
"""Genera MATERIAL_ENTRENAMIENTO_SESION_1.docx para el sistema EVA"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
AZUL_HUV   = RGBColor(0x1D, 0x29, 0x3D)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name, run.font.size, run.font.bold, run.font.italic = 'Calibri', Pt(size), bold, italic
    if color: run.font.color.rgb = color

def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size or sizes.get(level, 11), bold=True, color=color)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style, table.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        cell._tc.get_or_add_tcPr().append(shading)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F6FB' if ri % 2 == 0 else 'FFFFFF')
            cell._tc.get_or_add_tcPr().append(shading)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)
    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)
    doc.add_paragraph()

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin, section.right_margin = Cm(3), Cm(2.5)
    
    for _ in range(6): doc.add_paragraph()
    def cp(text, size=12, bold=False, color=GRIS_TEXTO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    cp("HOSPITAL UNIVERSITARIO DEL VALLE", 11, color=AZUL_HUV)
    cp("«Evaristo García» E.S.E.", 10, color=AZUL_HUV)
    doc.add_paragraph()
    cp("MATERIAL DE ENTRENAMIENTO — SESIÓN 1", 22, bold=True, color=AZUL_HUV)
    cp("Sistema EVA — Introducción y Módulo de Equipos", 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Versión 2.0.0  ·  Mayo 2026", 10, color=GRIS_TEXTO)
    doc.add_page_break()

    add_heading(doc, "1. Objetivos de la sesión")
    add_paragraph(doc, "Después de esta sesión, usted será capaz de:")
    add_bullet(doc, "Acceder correctamente al sistema EVA usando sus credenciales.")
    add_bullet(doc, "Navegar la interfaz y entender el propósito de cada módulo principal.")
    add_bullet(doc, "Consultar información de un equipo biomédico o industrial.")
    add_bullet(doc, "Realizar búsquedas avanzadas con filtros y criterios múltiples.")
    add_bullet(doc, "Descargar fichas técnicas en PDF y exportar datos a Excel.")

    add_heading(doc, "2. ¿Qué es EVA?")
    add_paragraph(doc, "EVA es el sistema centralizado de gestión de equipos médicos e industriales del HUV. Fue diseñado para reemplazar los procesos manuales anteriores y proporcionar:")
    add_bullet(doc, "Inventario digital unificado de más de 9.700 equipos.")
    add_bullet(doc, "Registro automático de mantenimientos preventivos y correctivos.")
    add_bullet(doc, "Control de calibraciones y vigencia INVIMA.")
    add_bullet(doc, "Sistema de tickets para órdenes de trabajo.")
    add_bullet(doc, "Reportes y estadísticas en tiempo real.")

    add_heading(doc, "3. Acceso al sistema")
    add_paragraph(doc, "Tabla 1. Instrucciones para acceder al sistema.")
    add_table(doc, ["Paso", "Acción"], [
        ["1", "Abra un navegador web (Chrome, Firefox, Safari, Edge)."],
        ["2", "Ingrese eva2.huv.gov.co en la barra de dirección."],
        ["3", "Ingrese su correo electrónico registrado en el sistema."],
        ["4", "Ingrese su contraseña personal."],
        ["5", "Haga clic en el botón 'Iniciar sesión'."],
        ["6", "Si ingresa datos incorrectos, verá un mensaje de error. Intente nuevamente."],
    ], col_widths=[2, 15])

    add_heading(doc, "4. Interfaz principal — Menú y navegación")
    add_paragraph(doc, "Tabla 2. Módulos principales accesibles desde el menú lateral.")
    add_table(doc, ["Módulo", "Descripción", "Quién lo usa"], [
        ["Dashboard",      "Panel de indicadores del sistema (KPIs y estadísticas).",                      "Administradores"],
        ["Equipos",        "Gestión del inventario de equipos biomédicos e industriales.",                  "Todo el personal"],
        ["Planes",         "Cronograma anual de mantenimiento preventivo.",                                 "Técnicos y coordinadores"],
        ["Órdenes",        "Sistema de tickets para correctivos y órdenes de trabajo.",                     "Todo el personal"],
        ["Repuestos",      "Inventario de repuestos disponibles.",                                          "Técnicos"],
        ["Calibraciones",  "Registro de calibraciones y control de vigencia.",                             "Coordinadores"],
        ["Capacitaciones", "Registro de entrenamientos del personal.",                                     "Coordinadores"],
        ["Administrador",  "Gestión de usuarios, roles, permisos y configuraciones avanzadas.",           "Super admins"],
    ], col_widths=[4, 8, 5])

    add_heading(doc, "5. Módulo de Equipos — Conceptos básicos")
    add_paragraph(doc, "El módulo de Equipos es el corazón del sistema. Mantiene la información completa de cada equipo: datos técnicos, ubicación, historial de mantenimientos, calibraciones, documentos adjuntos y registros INVIMA.")
    
    add_heading(doc, "5.1. Tipos de equipos", 2, AZUL_CLARO)
    add_bullet(doc, "Equipos biomédicos: Dispositivos médicos de uso clínico (monitores, ventiladores, desfibriladores, etc.).")
    add_bullet(doc, "Equipos industriales: Infraestructura de soporte (plantas eléctricas, compresores, bombas de agua, etc.).")

    add_heading(doc, "5.2. Consultar información de un equipo", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 3. Pasos para consultar un equipo.")
    add_table(doc, ["Paso", "Acción"], [
        ["1", "Haga clic en 'Equipos' en el menú lateral."],
        ["2", "Seleccione 'Biomédicos' o 'Industriales' según el tipo."],
        ["3", "Ingrese el nombre del equipo en el campo de búsqueda O use los filtros (servicio, estado, marca)."],
        ["4", "Haga clic en la fila del equipo que busca."],
        ["5", "Se abrirá un modal con toda la información: datos técnicos, mantenimientos, calibraciones, INVIMA, fotos."],
    ], col_widths=[2, 15])

    add_heading(doc, "5.3. Filtros avanzados", 2, AZUL_CLARO)
    add_paragraph(doc, "La tabla de equipos ofrece filtros para refinar búsquedas:")
    add_bullet(doc, "Por servicio (Urgencias, Quirófano, UTI, etc.).")
    add_bullet(doc, "Por estado (Activo, Baja, Préstamo, En reparación).")
    add_bullet(doc, "Por marca y modelo.")
    add_bullet(doc, "Por ubicación actual.")
    add_bullet(doc, "Texto libre (busca en nombre, código, descripción).")

    add_heading(doc, "6. Descargar documentos")
    add_paragraph(doc, "Tabla 4. Opciones de descarga disponibles para cada equipo.")
    add_table(doc, ["Documento", "Formato", "Contenido"], [
        ["Ficha técnica",   "PDF",   "Información completa del equipo en un solo documento imprimible."],
        ["Mantenimientos",  "Excel", "Historial de todos los mantenimientos preventivos y correctivos."],
        ["Calibraciones",   "Excel", "Registro completo de calibraciones realizadas."],
    ], col_widths=[5, 4, 9])

    add_heading(doc, "7. Práctica: Buscar y consultar un equipo")
    add_paragraph(doc, "Ejercicio práctico para consolidar lo aprendido:")
    add_bullet(doc, "Abra el módulo de Equipos → Biomédicos.")
    add_bullet(doc, "Busque un equipo del servicio 'URGENCIAS'.")
    add_bullet(doc, "Abra el detalle y revise el historial de mantenimientos.")
    add_bullet(doc, "Descargue la ficha técnica en PDF.")
    add_bullet(doc, "Vuelva a la lista y exporte los datos a Excel.")

    add_heading(doc, "8. Próxima sesión")
    add_paragraph(doc, "En la Sesión 2 aprenderemos a crear y gestionar tickets de mantenimiento correctivo, así como a registrar la ejecución del mantenimiento preventivo programado.")

    output_path = os.path.join(OUTPUT_DIR, "MATERIAL_ENTRENAMIENTO_SESION_1.docx")
    doc.save(output_path)
    print(f"Generado: {output_path}")

if __name__ == "__main__":
    build()
