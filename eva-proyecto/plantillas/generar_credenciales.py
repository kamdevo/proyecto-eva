# -*- coding: utf-8 -*-
"""
Genera CREDENCIALES_Y_CONFIGURACION.docx para el sistema EVA
Estilo: narrativo, con introducción, glosario, texto descriptivo y tablas.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

AZUL_HUV   = RGBColor(0x1D, 0x29, 0x3D)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)


def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name  = 'Calibri'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size or sizes.get(level, 11), bold=True, color=color)
    return p


def add_paragraph(doc, text, indent=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic, color=GRIS_TEXTO)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        cell._tc.tcPr.append(shading)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)

    for ri, row in enumerate(rows):
        bg = 'F2F6FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg)
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_pr.append(shading)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)

    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph()
    def cp(text, size=12, bold=False, color=GRIS_TEXTO, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
        return p
    cp("HOSPITAL UNIVERSITARIO DEL VALLE", 11, color=AZUL_HUV)
    cp("«Evaristo García» E.S.E.", 10, color=AZUL_HUV)
    doc.add_paragraph()
    cp("CREDENCIALES Y CONFIGURACIÓN", 22, bold=True, color=AZUL_HUV)
    cp("Sistema EVA", 16, bold=True, color=AZUL_CLARO)
    cp("Accesos, variables de entorno y configuración técnica", 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Versión 2.0.0  ·  Mayo 2026", 10, color=GRIS_TEXTO)
    cp("Documento confidencial — Prohibida su reproducción sin autorización.", 9, color=RGBColor(0x88, 0x88, 0x88))
    cp("Santiago de Cali, Colombia", 9, color=RGBColor(0x88, 0x88, 0x88))
    doc.add_page_break()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)
    add_cover(doc)

    add_heading(doc, "1. Introducción", 1, AZUL_HUV)
    add_paragraph(doc, "Este documento recopila todas las credenciales, URLs, variables de entorno y configuraciones técnicas necesarias para el despliegue, operación y soporte del sistema EVA en el Hospital Universitario del Valle. Está dirigido a administradores de sistemas, ingenieros biomédicos y personal de soporte técnico.")

    add_heading(doc, "2. Glosario de términos", 1, AZUL_CLARO)
    add_table(doc, ["Término", "Definición"], [
        ["URL base", "Dirección principal de acceso a la aplicación o API."],
        [".env", "Archivo de variables de entorno de Laravel o Vite."],
        ["Sanctum", "Sistema de autenticación de Laravel para APIs."],
        ["Token Bearer", "Token de autenticación enviado en el header Authorization."],
        ["Rol", "Nivel de acceso asignado a un usuario."],
        ["BD", "Base de datos relacional MySQL utilizada por el sistema."],
        ["Frontend", "Interfaz de usuario desarrollada en React."],
        ["Backend", "Servidor de lógica de negocio en Laravel."],
    ], col_widths=[4, 12])

    add_heading(doc, "3. Accesos y URLs principales", 1, AZUL_HUV)
    add_table(doc, ["Componente", "URL / Dirección", "Descripción"], [
        ["Frontend producción", "https://eva2.huv.gov.co", "Interfaz web para usuarios finales y admins."],
        ["Backend producción", "https://api.eva2.huv.gov.co", "API REST y lógica de negocio (Laravel)."],
        ["Frontend desarrollo", "http://localhost:5173", "Servidor local Vite (npm run dev)."],
        ["Backend desarrollo", "http://localhost:8000", "Servidor local Laravel (php artisan serve)."],
        ["Base de datos", "localhost:3306 (MySQL)", "Servidor de base de datos local o remoto."],
        ["Documentación técnica", "plantillas/DOCUMENTACION_TECNICA.docx", "Manual técnico completo del sistema."],
    ], col_widths=[4, 7, 7])

    add_heading(doc, "4. Variables de entorno principales (.env)", 1, AZUL_HUV)
    add_paragraph(doc, "Las siguientes variables deben configurarse en los archivos .env del backend y frontend para un funcionamiento correcto. No compartir estos valores fuera del equipo autorizado.")
    add_table(doc, ["Variable", "Ejemplo / Valor", "Descripción"], [
        ["APP_URL", "https://api.eva2.huv.gov.co", "URL base del backend Laravel."],
        ["DB_HOST", "localhost", "Host de la base de datos MySQL."],
        ["DB_DATABASE", "gestionthuv", "Nombre de la base de datos principal."],
        ["DB_USERNAME", "root", "Usuario de la base de datos."],
        ["DB_PASSWORD", "********", "Contraseña de la base de datos."],
        ["SANCTUM_STATEFUL_DOMAINS", "eva2.huv.gov.co,localhost:5173", "Dominios permitidos para autenticación API."],
        ["VITE_API_BASE_URL", "http://localhost:8000", "URL base de la API para el frontend (Vite)."],
        ["VITE_APP_NAME", "gestionthuv", "Nombre de la app en frontend."],
    ], col_widths=[4, 6, 8])

    add_heading(doc, "5. Roles y permisos de usuario", 1, AZUL_HUV)
    add_paragraph(doc, "El acceso a los módulos y funciones del sistema está controlado por roles y permisos. La siguiente tabla resume los roles principales:")
    add_table(doc, ["Rol", "ID", "Permisos principales"], [
        ["Super Administrador", "1", "Acceso total, gestión de usuarios, configuración avanzada."],
        ["Administrador", "2", "Acceso a todos los módulos operativos, sin gestión de super-admins."],
        ["Avanzado", "3", "Acceso a equipos, tickets, mantenimientos y calibraciones."],
        ["Normal", "4", "Solo lectura de equipos y gestión de sus propios tickets."],
    ], col_widths=[4, 2, 12])

    add_heading(doc, "6. Recomendaciones de seguridad", 1, AZUL_HUV)
    add_bullet(doc, "No compartir credenciales de base de datos ni archivos .env por correo o mensajería no segura.")
    add_bullet(doc, "Cambiar las contraseñas por defecto tras la instalación inicial.")
    add_bullet(doc, "Limitar el acceso SSH y MySQL solo a IPs autorizadas.")
    add_bullet(doc, "Mantener actualizado el sistema operativo y los paquetes de Laravel y React.")
    add_bullet(doc, "Realizar backups automáticos diarios de la base de datos y archivos adjuntos.")
    add_bullet(doc, "Revisar periódicamente los roles y permisos asignados a los usuarios.")

    output_path = os.path.join(OUTPUT_DIR, "CREDENCIALES_Y_CONFIGURACION.docx")
    doc.save(output_path)
    print(f"Generado: {output_path}")

if __name__ == "__main__":
    build()
