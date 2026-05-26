# -*- coding: utf-8 -*-
"""Genera MATERIAL_ENTRENAMIENTO_SESION_2.docx para el sistema EVA"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
AZUL_HUV   = RGBColor(0x1D, 0x29, 0x3D)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name, run.font.size, run.font.bold, run.font.italic = 'Calibri', Pt(size), bold, italic
    if color: run.font.color.rgb = color

def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size or sizes.get(level, 11), bold=True, color=color)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style, table.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        cell._tc.get_or_add_tcPr().append(shading)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F6FB' if ri % 2 == 0 else 'FFFFFF')
            cell._tc.get_or_add_tcPr().append(shading)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)
    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)
    doc.add_paragraph()

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin, section.right_margin = Cm(3), Cm(2.5)
    
    for _ in range(6): doc.add_paragraph()
    def cp(text, size=12, bold=False, color=GRIS_TEXTO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    cp("HOSPITAL UNIVERSITARIO DEL VALLE", 11, color=AZUL_HUV)
    cp("«Evaristo García» E.S.E.", 10, color=AZUL_HUV)
    doc.add_paragraph()
    cp("MATERIAL DE ENTRENAMIENTO — SESIÓN 2", 22, bold=True, color=AZUL_HUV)
    cp("Sistema EVA — Órdenes correctivas, mantenimiento preventivo y reportes", 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Versión 2.0.0  ·  Mayo 2026", 10, color=GRIS_TEXTO)
    doc.add_page_break()

    add_heading(doc, "1. Objetivos de la sesión")
    add_paragraph(doc, "Después de esta sesión, usted será capaz de:")
    add_bullet(doc, "Crear y gestionar tickets de mantenimiento correctivo desde el principio hasta el cierre.")
    add_bullet(doc, "Consultar y actualizar el estado de un ticket en ejecución.")
    add_bullet(doc, "Registrar la ejecución de mantenimientos preventivos programados.")
    add_bullet(doc, "Generar reportes personalizados en Excel con información histórica.")
    add_bullet(doc, "Entender los flujos de aprobación y los roles responsables en cada etapa.")

    add_heading(doc, "2. Módulo de Órdenes (Tickets)")
    add_paragraph(doc, "Los tickets representan las órdenes de trabajo para mantenimiento correctivo, reparación de emergencia u otros servicios no planeados.")

    add_heading(doc, "2.1. Estados de un ticket", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 1. Ciclo de vida de un ticket.")
    add_table(doc, ["Estado", "Descripción", "Quién puede cambiar"], [
        ["Abierto",      "El ticket fue creado pero aún no ha sido asignado a un técnico.",          "Coordinador"],
        ["En proceso",   "Un técnico está trabajando en la reparación.",                            "Técnico/Coordinador"],
        ["En espera",    "Se requieren repuestos o aprobación antes de continuar.",                 "Coordinador"],
        ["Resuelto",     "El técnico considera que la reparación está completa.",                  "Coordinador"],
        ["Verificado",   "El ingeniero responsable verificó que la solución es efectiva.",         "Coordinador"],
        ["Cerrado",      "El ticket fue cerrado, se archivó y se generó el reporte final.",       "Coordinador"],
    ], col_widths=[4.5, 8.5, 4.5])

    add_heading(doc, "2.2. Crear un nuevo ticket", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 2. Pasos para crear un ticket.")
    add_table(doc, ["Paso", "Acción"], [
        ["1", "Haga clic en 'Órdenes' → 'Mis Tickets' en el menú lateral."],
        ["2", "Haga clic en el botón 'Crear Ticket' (esquina superior derecha)."],
        ["3", "Seleccione el equipo afectado en el modal de búsqueda."],
        ["4", "Escriba una descripción clara del problema. Incluya cuándo empezó y qué síntomas observó."],
        ["5", "Si tiene fotografías o documentos, adjúntelos usando el botón 'Cargar archivo'."],
        ["6", "Seleccione el tipo de servicio (Emergencia, Reparación, Mantenimiento, Otro)."],
        ["7", "Haga clic en 'Crear'. El sistema generará un número único de ticket."],
    ], col_widths=[2, 15])

    add_heading(doc, "2.3. Actualizar un ticket en ejecución", 2, AZUL_CLARO)
    add_paragraph(doc, "Una vez asignado a un técnico, el ticket puede ser actualizado con avances:")
    add_bullet(doc, "El técnico registra diagnóstico inicial (qué cree que tiene el equipo).")
    add_bullet(doc, "El técnico adjunta evidencias fotográficas del estado actual.")
    add_bullet(doc, "El técnico cambia el estado a 'En proceso' cuando inicia la reparación.")
    add_bullet(doc, "El técnico adjunta fotos de la solución aplicada.")
    add_bullet(doc, "El técnico registra el tiempo dedicado (para análisis de productividad).")

    add_heading(doc, "2.4. Cerrar un ticket", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 3. Checklist antes de cerrar un ticket.")
    add_table(doc, ["Elemento", "Descripción"], [
        ["Diagnóstico completo",    "Está registrado y es claro."],
        ["Solución documentada",    "Incluye fotos, materiales usados, tiempo dedicado."],
        ["Verificación de función", "El equipo funciona correctamente después de la reparación."],
        ["Repuestos registrados",   "Todos los repuestos reemplazados están en el historial."],
        ["Observaciones finales",   "Se registraron recomendaciones para el futuro."],
        ["Firma o validación",      "El coordinador o ingeniero responsable verificó todo."],
    ], col_widths=[6.5, 11])

    add_heading(doc, "3. Módulo de Planes — Mantenimiento Preventivo")
    add_paragraph(doc, "El módulo de Planes permite consultar el cronograma anual de mantenimiento preventivo y registrar su ejecución.")

    add_heading(doc, "3.1. Consultar el cronograma anual", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 4. Estructura del cronograma de mantenimiento preventivo.")
    add_table(doc, ["Columna", "Información"], [
        ["Equipo",         "Nombre y código del equipo."],
        ["Servicio",       "Departamento donde está ubicado (Urgencias, UTI, Quirófano, etc.)."],
        ["Responsable",    "Técnico o coordinador responsable del mantenimiento."],
        ["Ene–Dic",        "Los 12 meses del año. Celda verde = mes asignado. Celda roja = se cumplió."],
        ["% Ejecución",    "Porcentaje de planes cumplidos en el año."],
        ["Observaciones",  "Notas sobre cambios, postergaciones o problemas."],
    ], col_widths=[4, 13])

    add_heading(doc, "3.2. Registrar ejecución de un plan", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 5. Pasos para registrar la ejecución de un plan.")
    add_table(doc, ["Paso", "Acción"], [
        ["1", "Localize el equipo y el mes programado en la tabla."],
        ["2", "Haga clic en el mes correspondiente."],
        ["3", "Se abrirá un modal para registrar la ejecución."],
        ["4", "Ingrese la fecha en que se realizó el mantenimiento."],
        ["5", "Agregue una descripción de lo que se hizo (limpieza, revisión, ajustes, etc.)."],
        ["6", "Adjunte evidencias fotográficas del proceso."],
        ["7", "Haga clic en 'Guardar'. El sistema marcará ese mes como cumplido."],
    ], col_widths=[2, 15])

    add_heading(doc, "4. Análisis de datos y reportes")
    add_paragraph(doc, "El sistema permite exportar datos en Excel para análisis adicional:")
    add_bullet(doc, "Exportar todos los equipos con su información técnica completa.")
    add_bullet(doc, "Exportar historial de correctivos con diagnóstico y tiempo de resolución.")
    add_bullet(doc, "Exportar cronograma con % de cumplimiento por equipo y por técnico.")
    add_bullet(doc, "Exportar calibraciones con fechas de vencimiento.")

    add_heading(doc, "5. Mejores prácticas para técnicos")
    add_bullet(doc, "Siempre registre diagnóstico claro: describe qué observó y qué fue lo encontrado.")
    add_bullet(doc, "Adjunte fotos de calidad: iluminación clara, ángulo que muestra el problema y la solución.")
    add_bullet(doc, "Registre tiempo real: no estime, anote cuántas horas realmente dedicó.")
    add_bullet(doc, "Documenta repuestos: ingrese marca, modelo y cantidad de cada repuesto usado.")
    add_bullet(doc, "Cierre rápido: no deje tickets pendientes más de 5 días sin actualizar estado.")

    add_heading(doc, "6. Solución de problemas")
    add_paragraph(doc, "Si encuentra algún problema al usar el sistema:")
    add_bullet(doc, "Tome una captura de pantalla del error.")
    add_bullet(doc, "Anote el paso que estaba realizando cuando ocurrió.")
    add_bullet(doc, "Cree un ticket de soporte incluiendo esta información.")
    add_bullet(doc, "El equipo de Bioingeniería lo contactará en los siguientes 24 horas.")

    add_heading(doc, "7. Próximas capacitaciones")
    add_paragraph(doc, "Temas avanzados para futuras sesiones (bajo demanda):")
    add_bullet(doc, "Módulo de Repuestos: Gestión de inventario y alertas de stock.")
    add_bullet(doc, "Módulo de Calibraciones: Control de vigencia INVIMA y reportes regulatorios.")
    add_bullet(doc, "Dashboard: Análisis de KPIs, tendencias y toma de decisiones.")
    add_bullet(doc, "Configuración de roles: Para administradores del sistema.")

    output_path = os.path.join(OUTPUT_DIR, "MATERIAL_ENTRENAMIENTO_SESION_2.docx")
    doc.save(output_path)
    print(f"Generado: {output_path}")

if __name__ == "__main__":
    build()
