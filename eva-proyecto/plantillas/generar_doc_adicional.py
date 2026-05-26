# -*- coding: utf-8 -*-
"""Genera DOCUMENTACION_TECNICA_ADICIONAL.docx para el sistema EVA"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
AZUL_HUV   = RGBColor(0x1D, 0x29, 0x3D)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB8)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
BLANCO     = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name, run.font.size, run.font.bold, run.font.italic = 'Calibri', Pt(size), bold, italic
    if color: run.font.color.rgb = color

def add_heading(doc, text, level=1, color=AZUL_HUV, size=None):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size or sizes.get(level, 11), bold=True, color=color)
    return p

def add_paragraph(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    if indent: p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10, color=GRIS_TEXTO)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style, table.alignment = 'Table Grid', WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1D293D')
        cell._tc.get_or_add_tcPr().append(shading)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9, bold=True, color=BLANCO)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F6FB' if ri % 2 == 0 else 'FFFFFF')
            cell._tc.get_or_add_tcPr().append(shading)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9, color=GRIS_TEXTO)
    if col_widths:
        for row_cells in table.rows:
            for i, width in enumerate(col_widths):
                row_cells.cells[i].width = Cm(width)
    doc.add_paragraph()

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin, section.right_margin = Cm(3), Cm(2.5)
    
    # Portada
    for _ in range(6): doc.add_paragraph()
    def cp(text, size=12, bold=False, color=GRIS_TEXTO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    cp("HOSPITAL UNIVERSITARIO DEL VALLE", 11, color=AZUL_HUV)
    cp("«Evaristo García» E.S.E.", 10, color=AZUL_HUV)
    doc.add_paragraph()
    cp("DOCUMENTACIÓN TÉCNICA AVANZADA", 22, bold=True, color=AZUL_HUV)
    cp("Sistema EVA — Módulos avanzados, flujos especiales y configuración", 10, color=GRIS_TEXTO)
    doc.add_paragraph()
    cp("Versión 2.0.0  ·  Mayo 2026", 10, color=GRIS_TEXTO)
    doc.add_page_break()

    add_heading(doc, "1. Introducción")
    add_paragraph(doc, "Este documento amplía la documentación técnica general del sistema EVA, cubriendo módulos avanzados, flujos especiales, configuraciones complejas y patrones de desarrollo utilizados en el sistema. Está dirigido a desarrolladores, arquitectos de sistemas e ingenieros biomédicos que requieren entender los detalles más profundos de la implementación.")

    add_heading(doc, "2. Estado del frontend — React Context y Hooks")
    add_paragraph(doc, "El frontend de EVA utiliza React Context para administrar el estado global de la aplicación, evitando prop drilling y facilitando el acceso a datos compartidos desde cualquier componente.")
    add_table(doc, ["Contexto", "Ubicación", "Responsabilidad"], [
        ["AuthContext",               "contexts/AuthContext.jsx",               "Usuario autenticado, token, rol, permisos, logout."],
        ["ToastContext",              "contexts/ToastContext.jsx",              "Notificaciones globales (éxito, error, advertencia)."],
        ["TicketsContext",            "contexts/TicketsContext.jsx",            "Estado compartido de tickets entre vistas."],
        ["EquipmentSearchContext",    "contexts/EquipmentSearchContext.jsx",    "Búsqueda global de equipos en la barra superior."],
    ], col_widths=[5, 5.5, 7])

    add_heading(doc, "3. Sistema de caché de equipmentPrefetchCache", 2, AZUL_CLARO)
    add_paragraph(doc, "El servicio equipmentPrefetchCache.js implementa un sistema de caché en memoria para las opciones de formularios. Evita peticiones repetidas al servidor cuando el usuario abre/cierra modales de equipos múltiples veces en la misma sesión.")
    add_bullet(doc, "Primer acceso al modal: GET /api/v1/equipos/filter-options carga todos los catálogos de una vez.")
    add_bullet(doc, "Accesos posteriores: Las opciones se obtienen del caché en memoria, sin peticiones HTTP.")
    add_bullet(doc, "Invalidación: El caché se limpia después de crear/editar un equipo, para reflejar cambios.")

    add_heading(doc, "4. Módulo de Calibraciones", 2, AZUL_CLARO)
    add_paragraph(doc, "Gestiona el registro y seguimiento de calibraciones de equipos biomédicos e industriales, con control automático de vigencia.")
    add_table(doc, ["Tabla BD", "Descripción", "Registros"], [
        ["calibracion",          "Calibraciones de equipos biomédicos",    "8.576"],
        ["calibracion_ind",      "Calibraciones de equipos industriales",  "321"],
    ], col_widths=[6, 7, 4])

    add_heading(doc, "5. Módulo de Repuestos", 2, AZUL_CLARO)
    add_paragraph(doc, "Inventario de repuestos disponibles para mantenimiento de equipos, con alertas automáticas de stock mínimo y trazabilidad de uso.")
    add_bullet(doc, "357 repuestos registrados en la base de datos.")
    add_bullet(doc, "Vinculación de repuestos pendientes a órdenes de trabajo correctivas.")
    add_bullet(doc, "Exportación de repuestos por servicio o por tipo de equipo.")

    add_heading(doc, "6. Flujo de mantenimiento correctivo", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 1. Flujo completo de un ticket correctivo desde su apertura hasta el cierre.")
    add_table(doc, ["Etapa", "Actor", "Acción"], [
        ["1. Creación",      "Personal clínico",     "Crea ticket indicando equipo y síntomas."],
        ["2. Diagnóstico",   "Ingeniero biomédico",  "Registra diagnóstico inicial y plan de acción."],
        ["3. Reparación",    "Técnico especialista", "Realiza mantenimiento y adjunta evidencias."],
        ["4. Verificación",  "Ingeniero responsable","Verifica que la solución es efectiva."],
        ["5. Cierre",        "Sistema / Usuario",    "Genera reporte PDF y archiva el ticket."],
    ], col_widths=[3, 4.5, 9.5])

    add_heading(doc, "7. Exportaciones complejas", 2, AZUL_CLARO)
    add_paragraph(doc, "El sistema permite exportar datos consolidados en formato Excel, con múltiples hojas y cálculos automáticos.")
    add_bullet(doc, "Exportación de cronograma anual con estado de ejecución por equipo.")
    add_bullet(doc, "Exportación de correctivos con diagnóstico, solución y tiempo de resolución.")
    add_bullet(doc, "Exportación de calibraciones con vigencia automática calculada.")

    add_heading(doc, "8. Seguridad avanzada", 2, AZUL_CLARO)
    add_paragraph(doc, "Tabla 2. Mecanismos avanzados de seguridad implementados.")
    add_table(doc, ["Mecanismo", "Descripción"], [
        ["Rate Limiting",              "Máximo 60 peticiones por minuto por IP."],
        ["CORS por dominio",           "Solo HUV y localhost (desarrollo) tienen acceso."],
        ["Validación strict de tipos", "Todas las peticiones tienen validación con FormRequest."],
        ["Encriptación de contraseña", "Bcrypt con 12 rondas (muy resistente a ataques)."],
        ["Timeout de sesión",          "Logout automático después de 30 minutos de inactividad."],
        ["Auditoría de cambios",       "Tabla cambios_cronograma registra todas las modificaciones con usuario y fecha."],
    ], col_widths=[6, 11])

    output_path = os.path.join(OUTPUT_DIR, "DOCUMENTACION_TECNICA_ADICIONAL.docx")
    doc.save(output_path)
    print(f"Generado: {output_path}")

if __name__ == "__main__":
    build()
